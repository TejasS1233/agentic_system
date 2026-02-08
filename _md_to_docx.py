"""Convert accomplished_tasks markdown to Word (.docx)."""
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

doc = Document()

# Style defaults
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10)

with open("docs/accomplished_tasks", "r", encoding="utf-8") as f:
    lines = f.readlines()

i = 0
while i < len(lines):
    line = lines[i].rstrip()
    
    # Skip empty lines
    if not line:
        i += 1
        continue
    
    # H1
    if line.startswith("# "):
        p = doc.add_heading(line[2:], level=0)
        i += 1
        continue
    
    # H2
    if line.startswith("## "):
        doc.add_heading(line[3:], level=1)
        i += 1
        continue
    
    # H3
    if line.startswith("### "):
        doc.add_heading(line[4:], level=2)
        i += 1
        continue
    
    # Blockquote
    if line.startswith(">"):
        p = doc.add_paragraph()
        p.style = 'Intense Quote' if 'Intense Quote' in [s.name for s in doc.styles] else 'Quote'
        p.add_run(line.lstrip("> "))
        i += 1
        continue
    
    # Horizontal rule
    if line.startswith("---"):
        p = doc.add_paragraph()
        p.add_run("─" * 60).font.color.rgb = RGBColor(180, 180, 180)
        i += 1
        continue
    
    # Table detection
    if "|" in line and i + 1 < len(lines) and "---" in lines[i + 1]:
        # Parse markdown table
        rows = []
        while i < len(lines) and "|" in lines[i]:
            cells = [c.strip() for c in lines[i].strip().split("|")]
            cells = [c for c in cells if c]  # Remove empty from leading/trailing |
            if not all(c.replace("-", "").replace(" ", "") == "" for c in cells):
                rows.append(cells)
            i += 1
        
        if rows:
            # Create table
            table = doc.add_table(rows=len(rows), cols=len(rows[0]))
            table.style = 'Light Grid Accent 1'
            for r_idx, row in enumerate(rows):
                for c_idx, cell_text in enumerate(row):
                    if c_idx < len(table.columns):
                        cell = table.cell(r_idx, c_idx)
                        cell.text = cell_text.replace("**", "")
                        if r_idx == 0:
                            for run in cell.paragraphs[0].runs:
                                run.bold = True
            doc.add_paragraph()  # spacing
        continue
    
    # Bullet / list item
    if line.startswith("- "):
        text = line[2:]
        p = doc.add_paragraph(style='List Bullet')
        # Handle **bold** segments
        parts = re.split(r'(\*\*.*?\*\*)', text)
        for part in parts:
            if part.startswith("**") and part.endswith("**"):
                run = p.add_run(part[2:-2])
                run.bold = True
            else:
                p.add_run(part)
        i += 1
        continue
    
    # Regular paragraph
    p = doc.add_paragraph()
    # Handle **bold**
    parts = re.split(r'(\*\*.*?\*\*)', line)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            run.bold = True
        else:
            p.add_run(part)
    i += 1

output = "docs/IASCIS_Accomplished_Tasks.docx"
doc.save(output)
print(f"Saved to: {output}")
